# Mode 2 便签笔记工作区

# ui/note_workspace.py
"""
Mode 2 - 便签笔记工作区（完整功能版）
支持实时自动保存、状态管理、重命名、上下移动、清空、导出等
"""

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QTextEdit,
    QGroupBox, QPushButton, QScrollArea, QFileDialog, QInputDialog,
    QMessageBox
)
from PyQt6.QtCore import Qt, QTimer
from PyQt6.QtGui import QFont, QTextCursor

from datetime import datetime
import os

from config import PYQT6_AVAILABLE,TIANGAN
from data_utils import save_notes
from time_utils import format_datetime
from .base_workspace import BaseWorkspace

# 尝试导入 python-docx（可选）
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

if not PYQT6_AVAILABLE:
    class NoteWorkspace(BaseWorkspace):
        def __init__(self, parent=None):
            super().__init__(parent)
else:
    class NoteWorkspace(BaseWorkspace):
        """便签笔记工作区"""
        def __init__(self, parent=None):
            super().__init__(parent)
            self.title_entry = None
            self.content_text = None
            self.log_label = None
            self.auto_save_timer = None
            self.edit_locked = True

        def build_ui(self):
            if self.ui_built:
                return
            self.ui_built = True

            layout = QVBoxLayout(self)
            layout.setContentsMargins(20, 20, 20, 10)
            layout.setSpacing(15)

            # 标题区
            title_layout = QHBoxLayout()
            title_label = QLabel("标题：")
            title_label.setFont(QFont("", 16, QFont.Weight.Bold))
            title_layout.addWidget(title_label)

            self.title_entry = QLineEdit()
            self.title_entry.setFont(QFont("", 16))
            self.title_entry.setPlaceholderText("输入便签标题...")
            self.title_entry.textChanged.connect(self.auto_save_draft)
            title_layout.addWidget(self.title_entry, stretch=1)
            layout.addLayout(title_layout)

            # 内容区
            content_label = QLabel("内容：")
            content_label.setFont(QFont("", 15, QFont.Weight.Bold))
            layout.addWidget(content_label)

            self.content_text = QTextEdit()
            self.content_text.setPlaceholderText("在这里写下你的想法、笔记...")
            self.content_text.setFont(QFont("", 13))
            self.content_text.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
            self.content_text.textChanged.connect(self.auto_save_draft)
            layout.addWidget(self.content_text, stretch=1)

            # 时间记录
            log_group = QGroupBox("时间记录")
            log_layout = QVBoxLayout()
            self.log_label = QLabel("创建时间：N/A\n更新时间：N/A")
            self.log_label.setStyleSheet("color: #666; font-size: 12px;")
            log_layout.addWidget(self.log_label)
            log_group.setLayout(log_layout)
            layout.addWidget(log_group)

            # note_workspace.py - build_ui() 末尾
            self.title_entry.setReadOnly(self.edit_locked)
            self.content_text.setReadOnly(self.edit_locked)

        def refresh_ui(self):
            main_window = self.window()
            if not hasattr(main_window, 'notes') or main_window.current_mode != 2:
                return

            idx = main_window.current_note_index
            note = main_window.notes[idx]

            # 更新标题和内容
            self.title_entry.setText(note.get("title", note.get("display_name", TIANGAN[main_window.current_note_index])))
            current_content = self.content_text.toPlainText().strip()
            saved_content = note.get("content", "").strip()
            if current_content != saved_content:
                self.content_text.setPlainText(saved_content)

            # 更新时间日志
            log_text = f"创建时间：{format_datetime(note.get('created_at', ''))}\n"
            log_text += f"更新时间：{format_datetime(note.get('updated_at', ''))}\n"
            if note.get("finished_at"):
                log_text += f"完成时间：{format_datetime(note.get('finished_at'))}"
            elif note.get("discarded_at"):
                log_text += f"废止时间：{format_datetime(note.get('discarded_at'))}"
            self.log_label.setText(log_text)

            # 根据状态控制编辑
            editable = note["status"] == "active"
            self.title_entry.setReadOnly(not editable)
            self.content_text.setReadOnly(not editable)

            # 可视反馈：完成/废止时标题变灰
            if not editable:
                self.title_entry.setStyleSheet("color: #888; font-style: italic;")
            else:
                self.title_entry.setStyleSheet("")

            # 每次刷新时，强制同步锁定状态
            self.title_entry.setReadOnly(self.edit_locked)
            self.content_text.setReadOnly(self.edit_locked)

        def auto_save_draft(self):
            """500ms 防抖自动保存"""
            if not self.auto_save_timer:
                self.auto_save_timer = QTimer(self)
                self.auto_save_timer.setSingleShot(True)
                self.auto_save_timer.timeout.connect(self._perform_auto_save)
            self.auto_save_timer.start(500)

        def _perform_auto_save(self):
            main_window = self.window()
            if not hasattr(main_window, 'notes'):
                return

            note = main_window.notes[main_window.current_note_index]
            new_title = self.title_entry.text().strip()
            new_content = self.content_text.toPlainText().strip()

            changed = False
            if new_title != note.get("title", "").strip():
                note["title"] = new_title
                changed = True
            if new_content != note.get("content", "").strip():
                note["content"] = new_content
                changed = True

            if changed:
                note["updated_at"] = datetime.now().isoformat()
                save_notes(main_window.notes)
                self.log_label.setText(
                    f"创建时间：{format_datetime(note.get('created_at', ''))}\n"
                    f"更新时间：{format_datetime(note['updated_at'])}"
                )
                main_window.statusBar().showMessage("便签已自动保存", 1500)

        # ===================== 通用操作实现 =====================
        def rename_current(self):
            main_window = self.window()
            idx = main_window.current_note_index
            note = main_window.notes[idx]
            default_name = TIANGAN[idx]  # 从 config 导入 TIANGAN
            current_name = note.get("display_name", default_name)
            new_name, ok = QInputDialog.getText(self, "重命名便签", "新名称：", text=current_name)
            if ok and new_name.strip():
                note["display_name"] = new_name.strip()
                save_notes(main_window.notes)
                main_window.refresh_left_list()

        def move_up_current(self):
            main_window = self.window()
            idx = main_window.current_note_index
            if idx > 0:
                main_window.notes[idx-1], main_window.notes[idx] = main_window.notes[idx], main_window.notes[idx-1]
                main_window.current_note_index = idx - 1
                save_notes(main_window.notes)
                main_window.refresh_left_list()

        def move_down_current(self):
            main_window = self.window()
            idx = main_window.current_note_index
            if idx < len(main_window.notes) - 1:
                main_window.notes[idx+1], main_window.notes[idx] = main_window.notes[idx], main_window.notes[idx+1]
                main_window.current_note_index = idx + 1
                save_notes(main_window.notes)
                main_window.refresh_left_list()

        def clear_current(self):
            main_window = self.window()
            idx = main_window.current_note_index
            note = main_window.notes[idx]
            if QMessageBox.question(self, "清空", "确定清空当前便签内容？") == QMessageBox.Yes:
                note["title"] = ""
                note["content"] = ""
                note["display_name"] = TIANGAN[idx]  # 恢复默认天干
                note["updated_at"] = datetime.now().isoformat()
                save_notes(main_window.notes)
                self.refresh_ui()
                main_window.refresh_left_list()
                main_window.statusBar().showMessage("当前便签已清空", 2000)

        # ===================== 编辑操作 =====================

        # note_workspace.py - 新增方法
        def lock_edit(self):
            self.edit_locked = True
            self.title_entry.setReadOnly(True)
            self.content_text.setReadOnly(True)
            self.window().statusBar().showMessage("已锁定编辑", 2000)

        def unlock_edit(self):
            note = self.window().notes[self.window().current_note_index]
            if note["status"] != "active":
                QMessageBox.information(self, "提示", "已完成或废止的便签无法编辑")
                return
            self.edit_locked = False
            self.title_entry.setReadOnly(False)
            self.content_text.setReadOnly(False)
            self.window().statusBar().showMessage("已解锁编辑，可编辑内容", 2000)

        # ===================== 状态管理 =====================
        def mark_complete(self):
            main_window = self.window()
            note = main_window.notes[main_window.current_note_index]
            if note["status"] != "active":
                return
            if QMessageBox.question(self, "标记完成", "确认将此便签标记为已完成？") == QMessageBox.Yes:
                note["status"] = "completed"
                note["finished_at"] = datetime.now().isoformat()
                note["updated_at"] = datetime.now().isoformat()
                save_notes(main_window.notes)
                self.refresh_ui()
                main_window.statusBar().showMessage("便签已标记为完成", 2000)

        def mark_discard(self):
            main_window = self.window()
            note = main_window.notes[main_window.current_note_index]
            if note["status"] != "active":
                return
            if QMessageBox.question(self, "标记废止", "确认将此便签标记为已废止？") == QMessageBox.Yes:
                note["status"] = "discarded"
                note["discarded_at"] = datetime.now().isoformat()
                note["updated_at"] = datetime.now().isoformat()
                save_notes(main_window.notes)
                self.refresh_ui()
                main_window.statusBar().showMessage("便签已标记为废止", 2000)

        # ===================== 导出 =====================
        def export_txt(self):
            main_window = self.window()
            note = main_window.notes[main_window.current_note_index]
            default_name = note.get("display_name", note.get("title", f"便签{main_window.current_note_index+1}"))
            path, _ = QFileDialog.getSaveFileName(self, "保存为 TXT", f"{default_name}.txt", "Text Files (*.txt)")
            if path:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(self._generate_export_content(note))
                main_window.statusBar().showMessage(f"已导出到 {os.path.basename(path)}", 3000)

        def export_as(self):
            main_window = self.window()
            note = main_window.notes[main_window.current_note_index]
            default_name = note.get("display_name", note.get("title", f"便签{main_window.current_note_index+1}"))
            path, filter_type = QFileDialog.getSaveFileName(
                self, "另存为", f"{default_name}",
                "Text Files (*.txt);;Word Files (*.docx)" if DOCX_AVAILABLE else "Text Files (*.txt)"
            )
            if not path:
                return
            content = self._generate_export_content(note)
            if path.endswith(".docx") and DOCX_AVAILABLE:
                doc = Document()
                doc.add_heading(note.get("display_name", "便签"), 0)
                doc.add_paragraph(f"状态：{'已完成' if note['status']=='completed' else '已废止' if note['status']=='discarded' else '进行中'}")
                doc.add_paragraph(f"创建：{format_datetime(note.get('created_at',''))}")
                doc.add_paragraph(f"更新：{format_datetime(note.get('updated_at',''))}")
                if note.get("finished_at"):
                    doc.add_paragraph(f"完成：{format_datetime(note['finished_at'])}")
                if note.get("discarded_at"):
                    doc.add_paragraph(f"废止：{format_datetime(note['discarded_at'])}")
                doc.add_paragraph("\n内容：\n" + content)
                doc.save(path)
            else:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)
            main_window.statusBar().showMessage(f"已保存到 {os.path.basename(path)}", 3000)

        def _generate_export_content(self, note):
            lines = [
                f"名称：{note.get('display_name', note.get('title', '未命名'))}",
                f"状态：{'已完成' if note['status']=='completed' else '已废止' if note['status']=='discarded' else '进行中'}",
                f"创建时间：{format_datetime(note.get('created_at', ''))}",
                f"更新时间：{format_datetime(note.get('updated_at', ''))}",
            ]
            if note.get("finished_at"):
                lines.append(f"完成时间：{format_datetime(note['finished_at'])}")
            if note.get("discarded_at"):
                lines.append(f"废止时间：{format_datetime(note['discarded_at'])}")
            lines.extend(["", "内容：", note.get("content", "")])
            return "\n".join(lines)